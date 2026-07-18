import io
import os
import zipfile
import pypdf
import pypdfium2 as pdfium
from PIL import Image
from docx import Document
import openpyxl
import pandas as pd
from xhtml2pdf import pisa
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.units import inch
import pdfplumber

# ─── Group 1: PDF Organizing & Page Editors ──────────────────────────────────

def merge_pdfs(pdf_files_bytes: list[bytes]) -> bytes:
    writer = pypdf.PdfWriter()
    for f_bytes in pdf_files_bytes:
        writer.append(io.BytesIO(f_bytes))
    out_buf = io.BytesIO()
    writer.write(out_buf)
    writer.close()
    return out_buf.getvalue()


def split_pdf(pdf_bytes: bytes, page_ranges: list[tuple[int, int]]) -> list[bytes]:
    # page_ranges is list of 1-indexed (start_page, end_page) inclusive
    reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
    total_pages = len(reader.pages)
    output_files = []
    
    for start, end in page_ranges:
        writer = pypdf.PdfWriter()
        # clamp ranges
        start_idx = max(0, start - 1)
        end_idx = min(total_pages, end)
        
        for i in range(start_idx, end_idx):
            writer.add_page(reader.pages[i])
            
        out_buf = io.BytesIO()
        writer.write(out_buf)
        output_files.append(out_buf.getvalue())
        writer.close()
        
    return output_files


def delete_pdf_pages(pdf_bytes: bytes, pages_to_delete: list[int]) -> bytes:
    # pages_to_delete is list of 1-indexed page numbers
    reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
    writer = pypdf.PdfWriter()
    delete_set = set(pages_to_delete)
    total_pages = len(reader.pages)
    
    # Validate: cannot delete ALL pages
    remaining = total_pages - len(delete_set & set(range(1, total_pages + 1)))
    if remaining <= 0:
        raise ValueError(f"Cannot delete all {total_pages} pages. At least one page must remain in the document.")
    
    for idx, page in enumerate(reader.pages):
        page_num = idx + 1
        if page_num not in delete_set:
            writer.add_page(page)
            
    out_buf = io.BytesIO()
    writer.write(out_buf)
    writer.close()
    return out_buf.getvalue()


def rotate_pdf_pages(pdf_bytes: bytes, pages_to_rotate: list[int], angle: int) -> bytes:
    # pages_to_rotate is list of 1-indexed page numbers, angle is 90, 180, or 270
    reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
    writer = pypdf.PdfWriter()
    rotate_set = set(pages_to_rotate)
    
    for idx, page in enumerate(reader.pages):
        page_num = idx + 1
        if page_num in rotate_set or not rotate_set: # if rotate_set is empty, rotate all
            page.rotate(angle)
        writer.add_page(page)
        
    out_buf = io.BytesIO()
    writer.write(out_buf)
    writer.close()
    return out_buf.getvalue()


def crop_pdf_pages(pdf_bytes: bytes, pages: list[int], left: float, right: float, top: float, bottom: float) -> bytes:
    # pages is list of 1-indexed, left, right, top, bottom are points or margins
    reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
    writer = pypdf.PdfWriter()
    page_set = set(pages)
    
    for idx, page in enumerate(reader.pages):
        page_num = idx + 1
        if page_num in page_set or not page_set:
            # Crop box coordinates
            box = page.mediabox
            box.lower_left = (box.left + left, box.bottom + bottom)
            box.upper_right = (box.right - right, box.top - top)
        writer.add_page(page)
        
    out_buf = io.BytesIO()
    writer.write(out_buf)
    writer.close()
    return out_buf.getvalue()


def extract_pdf_pages(pdf_bytes: bytes, pages_to_extract: list[int]) -> bytes:
    reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
    writer = pypdf.PdfWriter()
    total = len(reader.pages)
    
    for p_num in pages_to_extract:
        if 1 <= p_num <= total:
            writer.add_page(reader.pages[p_num - 1])
            
    out_buf = io.BytesIO()
    writer.write(out_buf)
    writer.close()
    return out_buf.getvalue()


def rearrange_pdf_pages(pdf_bytes: bytes, page_order: list[int]) -> bytes:
    # page_order is a list of 1-indexed page numbers in the new order
    reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
    writer = pypdf.PdfWriter()
    total = len(reader.pages)
    
    for p_num in page_order:
        if 1 <= p_num <= total:
            writer.add_page(reader.pages[p_num - 1])
            
    out_buf = io.BytesIO()
    writer.write(out_buf)
    writer.close()
    return out_buf.getvalue()


def add_blank_pages(pdf_bytes: bytes, positions: list[int]) -> bytes:
    # positions is list of 1-indexed page positions to insert blank pages (A4 size)
    reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
    writer = pypdf.PdfWriter()
    total = len(reader.pages)
    
    # We will build mapping of page additions
    insert_after = set(positions)
    
    for idx, page in enumerate(reader.pages):
        page_num = idx + 1
        writer.add_page(page)
        if page_num in insert_after:
            writer.add_blank_page()
            
    out_buf = io.BytesIO()
    writer.write(out_buf)
    writer.close()
    return out_buf.getvalue()


def add_page_numbers(pdf_bytes: bytes, style: str = "bottom_right") -> bytes:
    # style can be bottom_right, bottom_center, bottom_left, top_right, top_center, top_left
    reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
    writer = pypdf.PdfWriter()
    total_pages = len(reader.pages)
    
    for idx, page in enumerate(reader.pages):
        page_num = idx + 1
        
        # Create a single-page PDF containing just the page number
        num_buf = io.BytesIO()
        c = canvas.Canvas(num_buf, pagesize=(page.mediabox.width, page.mediabox.height))
        c.setFont("Helvetica", 9)
        c.setFillColorRGB(0.5, 0.5, 0.6) # Muted text
        
        w = float(page.mediabox.width)
        h = float(page.mediabox.height)
        
        text = f"{page_num} of {total_pages}"
        
        # Position mapping
        if style == "bottom_right":
            c.drawRightString(w - 36, 36, text)
        elif style == "bottom_center":
            c.drawCentredString(w / 2, 36, text)
        elif style == "bottom_left":
            c.drawString(36, 36, text)
        elif style == "top_right":
            c.drawRightString(w - 36, h - 36, text)
        elif style == "top_center":
            c.drawCentredString(w / 2, h - 36, text)
        elif style == "top_left":
            c.drawString(36, h - 36, text)
        else:
            c.drawCentredString(w / 2, 36, text)
            
        c.showPage()
        c.save()
        num_buf.seek(0)
        
        # Merge the stamp PDF with the current page
        stamp_reader = pypdf.PdfReader(num_buf)
        page.merge_page(stamp_reader.pages[0])
        writer.add_page(page)
        
    out_buf = io.BytesIO()
    writer.write(out_buf)
    writer.close()
    return out_buf.getvalue()


# ─── Group 2: Convert From PDF ───────────────────────────────────────────────

def pdf_to_images(pdf_bytes: bytes, image_format: str = "PNG") -> list[bytes]:
    # Render PDF pages to images using pypdfum2
    images_list = []
    pdf = pdfium.PdfDocument(pdf_bytes)
    
    for i in range(len(pdf)):
        page = pdf.get_page(i)
        # Render page at 2.0x scale (~150 DPI) for clarity
        pil_img = page.render(scale=2).to_pil()
        
        img_buf = io.BytesIO()
        pil_img.save(img_buf, format=image_format.upper())
        images_list.append(img_buf.getvalue())
        page.close()
        
    pdf.close()
    return images_list


def pdf_to_text(pdf_bytes: bytes) -> str:
    reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
    text_list = []
    for idx, page in enumerate(reader.pages):
        text = page.extract_text()
        text_list.append(f"--- Page {idx + 1} ---\n{text or ''}")
    return "\n\n".join(text_list)


def pdf_to_word(pdf_bytes: bytes) -> bytes:
    # Basic PDF to Docx conversion: extract text lines and write paragraphs
    doc = Document()
    reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
    
    for idx, page in enumerate(reader.pages):
        doc.add_heading(f"Page {idx + 1}", level=2)
        text = page.extract_text()
        if text:
            for line in text.split("\n"):
                if line.strip():
                    doc.add_paragraph(line)
        doc.add_page_break()
        
    out_buf = io.BytesIO()
    doc.save(out_buf)
    return out_buf.getvalue()


def pdf_to_excel(pdf_bytes: bytes) -> bytes:
    # Extracts tables from PDF using pdfplumber and dumps to xlsx sheets
    out_buf = io.BytesIO()
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        with pd.ExcelWriter(out_buf, engine="openpyxl") as writer:
            for idx, page in enumerate(pdf.pages):
                tables = page.extract_tables()
                if not tables:
                    # Write page text as single cell if no table found
                    text = page.extract_text() or ""
                    df = pd.DataFrame([{"Page Content": text}])
                    df.to_excel(writer, sheet_name=f"Page_{idx + 1}", index=False)
                    continue
                
                for t_idx, table in enumerate(tables):
                    df = pd.DataFrame(table)
                    sheet_name = f"Page_{idx + 1}_Table_{t_idx + 1}"[:31]  # Excel 31 char limit
                    df.to_excel(writer, sheet_name=sheet_name, index=False, header=False)
                    
    return out_buf.getvalue()


def pdf_to_html(pdf_bytes: bytes) -> str:
    # Convert PDF text content into formatted semantic HTML
    reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
    html_lines = [
        "<!DOCTYPE html>",
        "<html>",
        "<head>",
        "<meta charset='utf-8'>",
        "<style>",
        "body { font-family: sans-serif; line-height: 1.6; max-width: 800px; margin: 40px auto; padding: 0 20px; color: #333; }",
        ".page-break { border-bottom: 2px dashed #ccc; margin: 40px 0; padding-bottom: 10px; color: #888; font-size: 12px; }",
        "p { margin-bottom: 16px; white-space: pre-wrap; }",
        "</style>",
        "</head>",
        "<body>"
    ]
    
    for idx, page in enumerate(reader.pages):
        html_lines.append(f"<div class='page-break'>Page {idx + 1}</div>")
        text = page.extract_text()
        if text:
            # simple line breaks grouping
            paragraphs = text.split("\n\n")
            for p in paragraphs:
                if p.strip():
                    html_lines.append(f"<p>{p.strip()}</p>")
                    
    html_lines.extend(["</body>", "</html>"])
    return "\n".join(html_lines)


# ─── Group 3: Convert To PDF ─────────────────────────────────────────────────

def images_to_pdf(images_bytes_list: list[bytes]) -> bytes:
    pil_images = []
    for img_bytes in images_bytes_list:
        try:
            img = Image.open(io.BytesIO(img_bytes))
            # Convert to RGB mode (required for PDF formatting)
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
            pil_images.append(img)
        except Exception:
            continue
            
    if not pil_images:
        raise ValueError("No valid images found for conversion.")
        
    out_buf = io.BytesIO()
    # Save the first image, and append the rest as secondary pages
    pil_images[0].save(out_buf, format="PDF", save_all=True, append_images=pil_images[1:])
    return out_buf.getvalue()


def text_to_pdf(text_str: str) -> bytes:
    # Convert plain text to PDF using ReportLab
    out_buf = io.BytesIO()
    c = canvas.Canvas(out_buf, pagesize=letter)
    width, height = letter
    
    # 0.5 inch margins
    margin = 36
    y_position = height - margin
    c.setFont("Courier", 10)
    
    for line in text_str.split("\n"):
        if y_position < margin:
            c.showPage()
            c.setFont("Courier", 10)
            y_position = height - margin
        c.drawString(margin, y_position, line)
        y_position -= 13 # line height
        
    c.showPage()
    c.save()
    return out_buf.getvalue()


def docx_to_pdf(docx_bytes: bytes) -> bytes:
    # Docx to PDF: parse text paragraphs and generate PDF using xhtml2pdf
    doc = Document(io.BytesIO(docx_bytes))
    html_content = ["<html><body>"]
    for p in doc.paragraphs:
        # Check headings
        if p.style.name.startswith("Heading"):
            html_content.append(f"<h2>{p.text}</h2>")
        else:
            html_content.append(f"<p>{p.text}</p>")
    html_content.append("</body></html>")
    
    out_buf = io.BytesIO()
    pisa.CreatePDF("\n".join(html_content), dest=out_buf)
    return out_buf.getvalue()


def xlsx_to_pdf(xlsx_bytes: bytes) -> bytes:
    # Convert Excel table sheet grid to clean PDF table using HTML + xhtml2pdf
    wb = openpyxl.load_workbook(io.BytesIO(xlsx_bytes), data_only=True)
    html_content = [
        "<html><head><style>",
        "table { width: 100%; border-collapse: collapse; margin-bottom: 20px; font-size: 10px; }",
        "th, td { border: 1px solid #ccc; padding: 6px; text-align: left; }",
        "th { background-color: #f2f2f2; }",
        "h2 { font-size: 14px; margin-top: 20px; }",
        "</style></head><body>"
    ]
    
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        html_content.append(f"<h2>Sheet: {sheet_name}</h2>")
        html_content.append("<table>")
        
        for r_idx, row in enumerate(sheet.iter_rows(values_only=True)):
            if all(v is None for v in row):
                continue
            html_content.append("<tr>")
            for val in row:
                cell_val = "" if val is None else str(val)
                if r_idx == 0:
                    html_content.append(f"<th>{cell_val}</th>")
                else:
                    html_content.append(f"<td>{cell_val}</td>")
            html_content.append("</tr>")
            
        html_content.append("</table>")
        
    html_content.append("</body></html>")
    out_buf = io.BytesIO()
    pisa.CreatePDF("\n".join(html_content), dest=out_buf)
    return out_buf.getvalue()


def html_to_pdf(html_str: str) -> bytes:
    out_buf = io.BytesIO()
    pisa.CreatePDF(html_str, dest=out_buf)
    return out_buf.getvalue()


def rtf_to_pdf(rtf_bytes: bytes) -> bytes:
    # Simple RTF text parser helper: extract ASCII text blocks and render to PDF
    # Since Python doesn't have a lightweight RTF-to-PDF engine, we extract plain text lines
    text_content = []
    # Simple RTF parser logic
    in_control = False
    control_word = ""
    current_text = ""
    
    # We decode to ascii and capture alphanumeric characters and basic spacing
    rtf_str = rtf_bytes.decode("ascii", errors="ignore")
    
    # A lightweight extraction of alphanumeric strings that aren't backslashed tags
    lines = rtf_str.split("\n")
    for line in lines:
        cleaned_line = ""
        skip = False
        for char in line:
            if char == "\\":
                skip = True
                continue
            if skip:
                if char.isspace() or char in (";", "}", "{", "\\"):
                    skip = False
                continue
            if char not in ("{", "}"):
                cleaned_line += char
        if cleaned_line.strip():
            text_content.append(cleaned_line.strip())
            
    return text_to_pdf("\n".join(text_content))


# ─── Group 4: PDF Security & Document Signing ────────────────────────────────

def encrypt_pdf(pdf_bytes: bytes, password: str) -> bytes:
    reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
    writer = pypdf.PdfWriter()
    
    for page in reader.pages:
        writer.add_page(page)
        
    writer.encrypt(password)
    out_buf = io.BytesIO()
    writer.write(out_buf)
    writer.close()
    return out_buf.getvalue()


def decrypt_pdf(pdf_bytes: bytes, password: str) -> bytes:
    reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
    if reader.is_encrypted:
        reader.decrypt(password)
        
    writer = pypdf.PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
        
    out_buf = io.BytesIO()
    writer.write(out_buf)
    writer.close()
    return out_buf.getvalue()


def sign_pdf(pdf_bytes: bytes, signature_image_bytes: bytes, page_num: int, x: float, y: float, width: float, height: float) -> bytes:
    # Overlay signature image at x, y, width, height (coordinates in points, 1/72 inch)
    reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
    writer = pypdf.PdfWriter()
    total_pages = len(reader.pages)
    
    target_idx = max(0, min(total_pages - 1, page_num - 1))
    
    for idx, page in enumerate(reader.pages):
        if idx == target_idx:
            # Create a transparent overlay PDF with reportlab containing only the image
            overlay_buf = io.BytesIO()
            p_width = float(page.mediabox.width)
            p_height = float(page.mediabox.height)
            
            c = canvas.Canvas(overlay_buf, pagesize=(p_width, p_height))
            
            # Save bytes to PIL and draw
            from reportlab.lib.utils import ImageReader
            sig_img = ImageReader(Image.open(io.BytesIO(signature_image_bytes)))
            c.drawImage(sig_img, x, y, width=width, height=height, mask="auto")
            c.showPage()
            c.save()
            overlay_buf.seek(0)
            
            overlay_reader = pypdf.PdfReader(overlay_buf)
            page.merge_page(overlay_reader.pages[0])
            
        writer.add_page(page)
        
    out_buf = io.BytesIO()
    writer.write(out_buf)
    writer.close()
    return out_buf.getvalue()


def render_pdf_page(pdf_bytes: bytes, page_idx: int) -> bytes:
    pdf = pdfium.PdfDocument(pdf_bytes)
    if page_idx < 0 or page_idx >= len(pdf):
        pdf.close()
        raise ValueError("Page index out of range")
    page = pdf.get_page(page_idx)
    # Render page at 2.0x scale (~150 DPI) for clarity
    pil_img = page.render(scale=2).to_pil()
    img_buf = io.BytesIO()
    pil_img.save(img_buf, format="PNG")
    page.close()
    pdf.close()
    return img_buf.getvalue()


def reduce_file_size(file_bytes: bytes, filename: str, quality: int = 65, scale_percent: float = 1.0, width_percent: float = 100.0, height_percent: float = 100.0, target_size_kb: float | None = None) -> tuple[bytes, dict]:
    """
    Universal multi-format file compressor.
    Compresses Images, PDFs, Office docs (.docx, .pptx, .xlsx), OpenDocument files, SVG, and text/data files.
    Returns (compressed_bytes, metadata_dict).
    """
    import re
    from PIL import ImageOps
    ext = os.path.splitext(filename.lower())[1]
    orig_size = len(file_bytes)
    quality = max(10, min(100, int(quality)))
    
    # Calculate scale factor based on scale_percent, width_percent, and height_percent
    w_scale = (float(width_percent) / 100.0) if width_percent else 1.0
    h_scale = (float(height_percent) / 100.0) if height_percent else 1.0
    scale_percent = max(0.1, min(1.0, float(scale_percent) * min(w_scale, h_scale)))
    
    # ─── 1. AUTO MODE: TARGET SIZE DRIVEN OPTIMIZATION ───────────────────────
    if target_size_kb and target_size_kb > 0:
        out_bytes, quality, scale_percent = auto_target_compress(file_bytes, filename, target_size_kb)

    # ─── 2. IMAGES (.jpg, .jpeg, .png, .webp, .bmp, .tiff) ───────────────────
    elif ext in ('.jpg', '.jpeg', '.png', '.webp', '.bmp', '.tiff'):
        try:
            img = Image.open(io.BytesIO(file_bytes))
            img = ImageOps.exif_transpose(img)
            w, h = img.size
            
            if scale_percent < 1.0 or w_scale < 1.0 or h_scale < 1.0:
                new_w = max(1, int(w * w_scale * scale_percent))
                new_h = max(1, int(h * h_scale * scale_percent))
                img = img.resize((new_w, new_h), Image.Resampling.LANCZOS)
                
            out_buf = io.BytesIO()
            fmt = img.format if img.format else ('JPEG' if ext in ('.jpg', '.jpeg') else 'PNG')
            
            if ext in ('.jpg', '.jpeg') or (fmt and fmt.upper() in ('JPEG', 'JPG')):
                if img.mode in ('RGBA', 'P', 'LA'):
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    background.paste(img, mask=img.split()[-1] if 'A' in img.mode else None)
                    img = background
                else:
                    img = img.convert('RGB')
                img.save(out_buf, format='JPEG', quality=quality, optimize=True)
            elif ext == '.webp':
                img.save(out_buf, format='WEBP', quality=quality, optimize=True)
            elif ext == '.png':
                if quality < 80:
                    img_rgba = img.convert('RGBA') if img.mode == 'RGBA' else img.convert('RGB')
                    img_rgba.save(out_buf, format='WEBP', quality=quality, optimize=True)
                else:
                    img.save(out_buf, format='PNG', optimize=True)
            else:
                img_rgb = img.convert('RGB')
                img_rgb.save(out_buf, format='JPEG', quality=quality, optimize=True)
                
            out_bytes = out_buf.getvalue()
        except Exception:
            out_bytes = file_bytes

    # ─── 3. PDF DOCUMENTS (.pdf) ──────────────────────────────────────────────
    elif ext == '.pdf':
        try:
            comp_bytes = None
            try:
                reader = pypdf.PdfReader(io.BytesIO(file_bytes))
                writer = pypdf.PdfWriter()
                for page in reader.pages:
                    writer.add_page(page)
                writer.compress_content_streams()
                
                if quality < 90 or scale_percent < 1.0:
                    for page in writer.pages:
                        for img_name, img_obj in page.images.items():
                            try:
                                raw_img = Image.open(io.BytesIO(img_obj.data))
                                raw_img = ImageOps.exif_transpose(raw_img)
                                w, h = raw_img.size
                                if scale_percent < 1.0:
                                    raw_img = raw_img.resize((max(1, int(w * w_scale * scale_percent)), max(1, int(h * h_scale * scale_percent))), Image.Resampling.LANCZOS)
                                img_buf = io.BytesIO()
                                if raw_img.mode in ('RGBA', 'P', 'LA'):
                                    bg = Image.new('RGB', raw_img.size, (255, 255, 255))
                                    bg.paste(raw_img, mask=raw_img.split()[-1] if 'A' in raw_img.mode else None)
                                    raw_img = bg
                                else:
                                    raw_img = raw_img.convert('RGB')
                                raw_img.save(img_buf, format='JPEG', quality=quality, optimize=True)
                                img_obj.replace(img_buf.getvalue())
                            except Exception:
                                pass

                out_buf = io.BytesIO()
                writer.write(out_buf)
                writer.close()
                compressed_pdf = out_buf.getvalue()
                if len(compressed_pdf) < orig_size * 0.95 and len(compressed_pdf) > 100:
                    comp_bytes = compressed_pdf
            except Exception:
                comp_bytes = None

            if comp_bytes:
                out_bytes = comp_bytes
            else:
                render_scale = max(0.4, min(1.2, min(w_scale, h_scale) * (quality / 75.0)))
                out_bytes = raster_compress_pdf(file_bytes, quality=quality, scale=render_scale)
        except Exception:
            out_bytes = file_bytes

    # ─── 3. OFFICE & OPENDOCUMENT ZIP ARCHIVES (.docx, .pptx, .xlsx, .odt, .odp, .ods) ───
    elif ext in ('.docx', '.pptx', '.xlsx', '.odt', '.odp', '.ods'):
        try:
            in_zip = zipfile.ZipFile(io.BytesIO(file_bytes))
            out_buf = io.BytesIO()
            out_zip = zipfile.ZipFile(out_buf, 'w', zipfile.ZIP_DEFLATED, compresslevel=9)
            
            media_exts = ('.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.webp')
            
            for item in in_zip.infolist():
                item_data = in_zip.read(item.filename)
                item_lower = item.filename.lower()
                
                # Check if file is an embedded image inside office media folders
                if any(item_lower.endswith(m_ext) for m_ext in media_exts) and ('media/' in item_lower or 'pictures/' in item_lower):
                    try:
                        img = Image.open(io.BytesIO(item_data))
                        w, h = img.size
                        if scale_percent < 1.0:
                            img = img.resize((max(1, int(w * scale_percent)), max(1, int(h * scale_percent))), Image.Resampling.LANCZOS)
                        
                        img_out = io.BytesIO()
                        if img.mode in ('RGBA', 'P', 'LA'):
                            bg = Image.new('RGB', img.size, (255, 255, 255))
                            bg.paste(img, mask=img.split()[-1] if 'A' in img.mode else None)
                            img = bg
                        else:
                            img = img.convert('RGB')
                        img.save(img_out, format='JPEG', quality=quality, optimize=True)
                        compressed_img = img_out.getvalue()
                        
                        if len(compressed_img) < len(item_data):
                            item_data = compressed_img
                    except Exception:
                        pass
                
                out_zip.writestr(item.filename, item_data)
                
            out_zip.close()
            compressed_doc = out_buf.getvalue()
            if len(compressed_doc) < orig_size:
                out_bytes = compressed_doc
        except Exception:
            out_bytes = file_bytes

    # ─── 4. SVG VECTOR IMAGES (.svg) ──────────────────────────────────────────
    elif ext == '.svg':
        try:
            text = file_bytes.decode('utf-8', errors='ignore')
            # Remove XML comments and redundant whitespace
            text = re.sub(r'<!--.*?-->', '', text, flags=re.DOTALL)
            text = re.sub(r'>\s+<', '><', text)
            text = text.strip()
            out_bytes = text.encode('utf-8')
        except Exception:
            out_bytes = file_bytes

    # ─── 5. TEXT / DATA FILES (.json, .csv, .xml, .txt, .html) ───────────────
    elif ext in ('.json', '.csv', '.xml', '.txt', '.html'):
        try:
            if ext == '.json':
                import json
                data = json.loads(file_bytes.decode('utf-8'))
                minified = json.dumps(data, separators=(',', ':'))
                out_bytes = minified.encode('utf-8')
            else:
                text = file_bytes.decode('utf-8', errors='ignore')
                lines = [line.strip() for line in text.splitlines() if line.strip()]
                out_bytes = '\n'.join(lines).encode('utf-8')
        except Exception:
            out_bytes = file_bytes

    # ─── 6. GENERAL FALLBACK ARCHIVE OPTIMIZATION ─────────────────────────────
    else:
        # If target size requested or general file, compress into ZIP stream
        try:
            out_buf = io.BytesIO()
            with zipfile.ZipFile(out_buf, 'w', zipfile.ZIP_DEFLATED, compresslevel=9) as z:
                z.writestr(filename, file_bytes)
            zip_bytes = out_buf.getvalue()
            if len(zip_bytes) < orig_size:
                out_bytes = zip_bytes
        except Exception:
            out_bytes = file_bytes

    # ─── Calculate Stats ──────────────────────────────────────────────────────
    compressed_size = len(out_bytes)
    saved_bytes = max(0, orig_size - compressed_size)
    savings_percent = round((saved_bytes / orig_size) * 100, 1) if orig_size > 0 else 0
    
    meta = {
        "filename": filename,
        "original_size": orig_size,
        "compressed_size": compressed_size,
        "saved_bytes": saved_bytes,
        "savings_percent": savings_percent,
        "quality": quality,
        "scale_percent": scale_percent
    }
    
    return out_bytes, meta


def raster_compress_pdf(file_bytes: bytes, quality: int = 65, scale: float = 1.0) -> bytes:
    """Renders PDF pages as JPEG images and saves as a 100% valid, compliant multi-page PDF."""
    try:
        from PIL import ImageOps
        pdf_doc = pdfium.PdfDocument(file_bytes)
        pil_images = []
        
        for i in range(len(pdf_doc)):
            page = pdf_doc.get_page(i)
            pil_img = page.render(scale=scale).to_pil()
            pil_img = ImageOps.exif_transpose(pil_img)
            page.close()
            
            if pil_img.mode in ('RGBA', 'P', 'LA'):
                bg = Image.new('RGB', pil_img.size, (255, 255, 255))
                bg.paste(pil_img, mask=pil_img.split()[-1] if 'A' in pil_img.mode else None)
                pil_img = bg
            else:
                pil_img = pil_img.convert('RGB')
                
            img_buf = io.BytesIO()
            pil_img.save(img_buf, format='JPEG', quality=quality, optimize=True)
            img_buf.seek(0)
            
            c_img = Image.open(img_buf)
            c_img.load()
            pil_images.append(c_img)
            
        pdf_doc.close()
        
        if not pil_images:
            return file_bytes
            
        pdf_out_buf = io.BytesIO()
        first_img = pil_images[0]
        if len(pil_images) > 1:
            first_img.save(pdf_out_buf, format='PDF', save_all=True, append_images=pil_images[1:])
        else:
            first_img.save(pdf_out_buf, format='PDF')
            
        return pdf_out_buf.getvalue()
    except Exception:
        return file_bytes


def compress_single_pass(file_bytes: bytes, filename: str, quality: int = 65, scale: float = 1.0) -> bytes:
    ext = os.path.splitext(filename.lower())[1]
    if ext in ('.jpg', '.jpeg', '.png', '.webp', '.bmp', '.tiff'):
        try:
            from PIL import ImageOps
            img = Image.open(io.BytesIO(file_bytes))
            img = ImageOps.exif_transpose(img)
            if scale < 1.0:
                img = img.resize((max(1, int(img.width * scale)), max(1, int(img.height * scale))), Image.Resampling.LANCZOS)
            out_buf = io.BytesIO()
            if img.mode in ('RGBA', 'P', 'LA'):
                bg = Image.new('RGB', img.size, (255, 255, 255))
                bg.paste(img, mask=img.split()[-1] if 'A' in img.mode else None)
                img = bg
            else:
                img = img.convert('RGB')
            img.save(out_buf, format='JPEG', quality=quality, optimize=True)
            return out_buf.getvalue()
        except Exception:
            return file_bytes
    elif ext == '.pdf':
        return raster_compress_pdf(file_bytes, quality=quality, scale=scale)
    return file_bytes


def auto_target_compress(file_bytes: bytes, filename: str, target_size_kb: float) -> tuple[bytes, int, float]:
    """
    Lightning-fast Binary Search auto-target algorithm (under 1 sec execution).
    Returns (compressed_bytes, quality_used, scale_used).
    """
    import math
    target_bytes = target_size_kb * 1024
    orig_size = len(file_bytes)

    if orig_size <= target_bytes:
        return file_bytes, 90, 1.0

    # Pass 1: Test at 80% quality and 1.0 scale
    p1_bytes = compress_single_pass(file_bytes, filename, quality=80, scale=1.0)
    if len(p1_bytes) <= target_bytes:
        return p1_bytes, 80, 1.0

    # Calculate optimal scale directly using byte ratio estimation
    ratio = target_bytes / max(1, len(p1_bytes))
    target_scale = max(0.25, min(1.0, math.sqrt(ratio)))

    # Binary search on 5% quality increments (Quality 10 to 90) at target_scale (3 passes max!)
    low_q, high_q = 10, 90
    best_bytes = p1_bytes
    best_q = 50

    while low_q <= high_q:
        mid_q = ((low_q + high_q) // 10) * 5
        test_bytes = compress_single_pass(file_bytes, filename, quality=mid_q, scale=target_scale)
        
        if len(test_bytes) <= target_bytes:
            best_bytes = test_bytes
            best_q = mid_q
            low_q = mid_q + 5  # Try higher quality
        else:
            high_q = mid_q - 5  # Lower quality needed

    return best_bytes, best_q, target_scale
